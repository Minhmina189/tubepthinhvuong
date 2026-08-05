# -*- coding: utf-8 -*-
"""
Multi-format input parser for Nam An quotation system.
Excel/CSV: smart column detection, no AI required.
Word/PDF/Images: Claude Vision API (requires ANTHROPIC_API_KEY).
"""
import base64
import json
import re
import tempfile
from datetime import datetime as _dt
from pathlib import Path
from typing import List, Optional


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────

def parse_input_file(filepath: str) -> List[dict]:
    """
    Universal parser for any supported input format.
    Returns list[dict] with keys: name, material, qty, note_in
    """
    ext = Path(filepath).suffix.lower()

    if ext in {".xlsx", ".xls"}:
        return _smart_parse_excel(filepath)

    if ext == ".csv":
        return _smart_parse_csv(filepath)

    if ext in {".docx", ".doc"}:
        text = _extract_text_docx(filepath)
        return _ai_extract_from_text(text)

    if ext == ".pdf":
        text = _extract_text_pdf(filepath)
        if text and len(text.strip()) > 50:
            return _ai_extract_from_text(text)
        image_paths = _pdf_to_images(filepath)
        return _ai_extract_from_images(image_paths)

    if ext in {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}:
        return _ai_extract_from_images([filepath])

    raise ValueError(f"Định dạng file không được hỗ trợ: {ext}")


# ─────────────────────────────────────────────
# DIN / ISO NOTATION TABLES
# ─────────────────────────────────────────────

# DIN 125 flat washer: hole diameter (mm) → bolt size
_DIN125_HOLE_BOLT = {
    '3.2': 3,  '4.3': 4,  '5.3': 5,  '6.4': 6,
    '8.4': 8,  '10.5': 10, '13': 12, '15': 14,
    '17': 16,  '19': 18,  '21': 20,  '23': 22,
    '25': 24,  '28': 27,  '31': 30,  '34': 33, '37': 36,
}

# DIN mapping loaded from din_mapping.json (editable via /din-mapping UI)
_DIN_MAP_FILE = Path(__file__).parent / "din_mapping.json"
_DIN_MAP_CACHE: dict = {}
_DIN_MAP_MTIME: float = 0.0


def _get_din_map() -> dict:
    """Load din_mapping.json, reloading if the file changed."""
    global _DIN_MAP_CACHE, _DIN_MAP_MTIME
    try:
        mtime = _DIN_MAP_FILE.stat().st_mtime
        if mtime != _DIN_MAP_MTIME:
            raw = json.loads(_DIN_MAP_FILE.read_text(encoding='utf-8'))
            _DIN_MAP_CACHE = {k: v for k, v in raw.items() if not k.startswith('_')}
            _DIN_MAP_MTIME = mtime
    except Exception:
        pass
    return _DIN_MAP_CACHE


# Fallback table (used if din_mapping.json is missing)
_DIN_PTYPE_FALLBACK = {
    '125': 'Đệm phẳng',   '126': 'Đệm phẳng',
    '127': 'Đệm vênh',    '128': 'Đệm vênh',
    '933': 'Bu lông',      '931': 'Bu lông',
    '960': 'Bu lông',      '961': 'Bu lông',
    '912': 'Lục giác chìm trụ',
    '7380': 'Lục giác chìm tròn',
    '934': 'Đai ốc lục giác', '985': 'Đai ốc hãm',
    '1587': 'Đai ốc mũ chóp',
    '94': 'Chốt chẻ', '7337': 'Đinh rút',
}


def _din_info(din_num: str) -> dict:
    """Return the mapping entry for a DIN number."""
    m = _get_din_map()
    if din_num in m:
        return m[din_num]
    # fallback
    ten = _DIN_PTYPE_FALLBACK.get(din_num, '')
    return {'ten': ten, 'co_chieu_dai': ten in ('Bu lông', 'Lục giác chìm trụ',
                                                  'Lục giác chìm tròn', 'Lục giác chìm đầu bằng',
                                                  'Chốt chẻ', 'Đinh rút'),
            'ghi_chu': '', 'ky_hieu': ''}


# ─────────────────────────────────────────────
# SMART EXCEL PARSER (no AI)
# ─────────────────────────────────────────────

# Role → keyword → score  (higher score = stronger match)
# Include both Vietnamese (with diacritics) and ASCII-fallback variants
_COL_ROLES = {
    'name': {
        # Vietnamese variants
        'tên chủng loại': 30, 'tên hàng hóa': 28, 'tên hàng hoá': 28,
        'tên vật tư': 28, 'tên sản phẩm': 25,
        'hàng hóa': 20, 'hàng hoá': 20, 'vật tư': 15,
        'sản phẩm': 15, 'tên hàng': 15, 'mô tả': 10, 'diễn giải': 10,
        # ASCII / English fallbacks
        'ten chung loai': 28, 'ten hang hoa': 25, 'ten vat tu': 25,
        'ten san pham': 22, 'hang hoa': 18, 'vat tu': 12,
        'description': 10, 'product': 10, 'item': 8, 'ten': 4,
    },
    'simplified_name': {            # e.g. "Ghi chú kiểm tra" column
        'ghi chú kiểm tra': 100, 'ghi chú kt': 100, 'ghi chú ktra': 100,
        'ghi chu kiem tra': 100, 'ghi chu kt': 100,
    },
    'qty': {
        'số lượng mua': 35, 'số lượng': 25, 'so luong mua': 33, 'so luong': 23,
        'quantity': 20, 'qty': 20, 'sl': 15,
    },
    'material': {
        'vật liệu chính': 35, 'chất liệu chính': 35,
        'vật liệu': 25, 'chất liệu': 25,
        'vat lieu chinh': 33, 'chat lieu chinh': 33,
        'vat lieu': 23, 'chat lieu': 23,
        'material': 20, 'chất': 8,
    },
    'note': {
        'ghi chú': 10, 'ghi chu': 9, 'note': 8, 'chú thích': 8,
    },
}

# Regex: does the cell value look like a product?
_PRODUCT_RE = re.compile(
    r'(bu\s*l[ôo]ng|bulông|đ[eệ]m\s*(ph[aằ]ng|v[eê]nh)|ecu'
    r'|đai\s*[oóô]c|l[uụ]c\s*gi[aá]c|v[ií]t|ch[oô]t\s*ch[eẻ]'
    r'|đ[iì]nh\s*r[uú]t|thanh\s*ren|\bren\b|m\d+)',
    re.IGNORECASE
)


def _smart_parse_excel(filepath: str) -> List[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    all_items: List[dict] = []

    for ws in wb.worksheets:
        rows = [tuple(row) for row in ws.iter_rows(values_only=True)]
        if not rows:
            continue
        header_idx, col_map = _detect_columns(rows)
        if header_idx is None or 'name' not in col_map:
            continue
        for row in rows[header_idx + 1:]:
            if not any(v is not None for v in row):
                continue
            item = _row_to_item(row, col_map)
            if item:
                all_items.append(item)

    return all_items


def _smart_parse_csv(filepath: str) -> List[dict]:
    import csv
    rows: List[tuple] = []
    for enc in ['utf-8-sig', 'utf-8', 'cp1252', 'latin-1']:
        try:
            with open(filepath, encoding=enc, newline='') as f:
                rows = [tuple(v.strip() or None for v in r) for r in csv.reader(f)]
            break
        except (UnicodeDecodeError, Exception):
            continue
    if not rows:
        return []

    header_idx, col_map = _detect_columns(rows)
    if header_idx is None or 'name' not in col_map:
        return []

    items: List[dict] = []
    for row in rows[header_idx + 1:]:
        if not any(v for v in row):
            continue
        item = _row_to_item(row, col_map)
        if item:
            items.append(item)
    return items


def _detect_columns(rows: list):
    """Find best header row and return (header_idx, col_map)."""
    best_score, best_idx, best_map = 0, None, {}
    for i, row in enumerate(rows[:25]):
        score, col_map = _score_header_row(row)
        if score > best_score and 'name' in col_map:
            best_score, best_idx, best_map = score, i, col_map
    return best_idx, best_map


def _score_header_row(row: tuple):
    """Return (total_score, col_map) treating this row as a header."""
    # Step 1: for each cell, find the best matching role and its score
    cell_role_scores: dict = {}   # col_idx → {role: score}
    for col_idx, val in enumerate(row):
        if val is None:
            continue
        cell = str(val).lower().strip()
        if not cell:
            continue
        for role, keywords in _COL_ROLES.items():
            for kw, score in keywords.items():
                if kw in cell:
                    cell_role_scores.setdefault(col_idx, {})[role] = max(
                        cell_role_scores.get(col_idx, {}).get(role, 0), score
                    )
                    # don't break — allow multiple roles per cell, resolve below

    # Step 2: each column gets the role with the highest score
    col_best: dict = {}   # col_idx → (role, score)
    for col_idx, role_scores in cell_role_scores.items():
        best_role = max(role_scores, key=role_scores.get)
        col_best[col_idx] = (best_role, role_scores[best_role])

    # Step 3: each role gets the column with the highest score
    role_col: dict = {}   # role → (col_idx, score)
    for col_idx, (role, score) in col_best.items():
        if role not in role_col or score > role_col[role][1]:
            role_col[role] = (col_idx, score)

    col_map = {role: col_idx for role, (col_idx, _) in role_col.items()}
    total = sum(s for _, s in role_col.values())
    return total, col_map


def _row_to_item(row: tuple, col_map: dict) -> Optional[dict]:
    """Extract one product item from a data row."""
    def get(key: str) -> str:
        idx = col_map.get(key)
        if idx is None or idx >= len(row):
            return ''
        v = row[idx]
        if v is None:
            return ''
        if isinstance(v, _dt):     # skip date cells
            return ''
        return str(v).strip()

    raw_name = get('name')
    simplified = get('simplified_name')

    # Prefer simplified name (e.g. "Ghi chú kiểm tra") if it has M-size notation
    if simplified and re.search(r'\bM\d+', simplified, re.IGNORECASE):
        name = simplified
    elif raw_name:
        name = _normalize_din_name(raw_name)
    else:
        return None

    if not name or not _PRODUCT_RE.search(name):
        return None

    # Quantity
    qty = 1
    qty_raw = get('qty')
    if qty_raw:
        try:
            qty = max(1, int(float(qty_raw.replace(',', '.'))))
        except (ValueError, TypeError):
            qty = 1

    return {
        'name':      name,
        'material':  _normalize_material(get('material')),
        'qty':       qty,
        'note_in':   get('note'),
    }


# ─────────────────────────────────────────────
# NAME & MATERIAL NORMALIZERS
# ─────────────────────────────────────────────

def _normalize_din_name(raw: str) -> str:
    """Convert DIN/ISO notation to Nam An standard name using din_mapping.json."""
    s = raw.strip()

    # Extract DIN / ISO standard number
    din_m = re.search(r'\bDIN\s*(\d+)', s, re.IGNORECASE)
    iso_m = re.search(r'\bISO\s*(\d+)', s, re.IGNORECASE)
    std_num = din_m or iso_m
    din_num = std_num.group(1) if std_num else None

    # Pass-through: already clean Vietnamese name (no DIN code present)
    _CLEAN_KW = re.compile(
        r'^(bu\s*l[ôo]ng|đ[eệ]m|ecu|đai\s*[oóô]c|l[uụ]c\s*gi[aá]c|v[ií]t'
        r'|ch[oô]t\s*ch[eẻ]|đ[iì]nh\s*r[uú]t|thanh\s*ren)',
        re.IGNORECASE
    )
    if not std_num and _CLEAN_KW.match(s):
        return s

    # Look up DIN info from user's mapping table
    info = _din_info(din_num) if din_num else {}
    ten  = info.get('ten', '')
    has_len = info.get('co_chieu_dai', False)
    ky_hieu = info.get('ky_hieu', '')   # 'D' for chốt chẻ, 'A' for đinh rút

    # ── Flat washer  D{hole} - DIN 125/126 ──────────────────────
    if din_num in ('125', '126') or ten == 'Đệm phẳng':
        m = re.search(r'\bD(\d+\.?\d*)\b', s)
        if m:
            bolt = _DIN125_HOLE_BOLT.get(m.group(1))
            label = ten or 'Đệm phẳng'
            return f"{label} M{bolt}" if bolt else f"{label} D{m.group(1)}"

    # ── Spring washer  A/B{size} - DIN 127/128 ───────────────────
    if din_num in ('127', '128') or ten == 'Đệm vênh':
        m = re.search(r'\b[AB](\d+)\b', s)
        if m:
            return f"{ten or 'Đệm vênh'} M{m.group(1)}"

    # ── D-notation parts (Chốt chẻ…): D{dia}xL{len} ─────────────
    if ky_hieu == 'D':
        m = re.search(r'\bD([\d.]+)xL(\d+)', s, re.IGNORECASE)
        if m:
            return f"{ten} D{m.group(1)}x{m.group(2)}"

    # ── A-notation parts (Đinh rút…): A{dia}x{len} ───────────────
    if ky_hieu == 'A':
        m = re.search(r'\bA([\d.]+)x(\d+)', s, re.IGNORECASE)
        if m:
            return f"{ten} A{m.group(1)}x{m.group(2)}"

    # ── Parts without length (nuts, washers via fallback) ─────────
    if ten and not has_len:
        m = re.search(r'\bM(\d+)', s, re.IGNORECASE)
        if m:
            return f"{ten} M{m.group(1)}"

    # ── Bolt/screw with pitch notation: M{size}x{pitch}xL{length} ─
    m = re.search(r'\bM(\d+)x[\d.]+xL(\d+)', s, re.IGNORECASE)
    if m:
        sz, ln = m.group(1), m.group(2)
        return f"{ten or 'Bu lông'} M{sz}x{ln}"

    # ── Bolt/screw: M{size}x{length} (2+ digits to avoid matching pitch) ─
    m = re.search(r'\bM(\d+)x(\d{2,})', s, re.IGNORECASE)
    if m:
        sz, ln = m.group(1), m.group(2)
        return f"{ten or 'Bu lông'} M{sz}x{ln}"

    # ── Fallback: just M{size} ────────────────────────────────────
    m = re.search(r'\bM(\d+)', s, re.IGNORECASE)
    if m and ten:
        return f"{ten} M{m.group(1)}"

    return s   # unknown format — pass through


def _normalize_material(raw: str) -> str:
    """Map verbose material descriptions to Nam An standard keywords."""
    if not raw:
        return ''
    sl = raw.lower()

    # Use specific patterns to avoid false matches (e.g. "201" inside "2012")
    if re.search(r'\binox\s*316|\bss\s*316',           sl): return 'Inox 316'
    if re.search(r'\binox\s*201|\bss\s*201',           sl): return 'Inox 201'
    if re.search(r'\binox\b|\b304\b',                  sl): return 'Inox 304'
    # Mạ kẽm — phải check TRƯỚC cấp độ bền để kẽm không bị nuốt bởi "8.8"/"10.9"
    _MA_KEM_PAT = r'mạ\s*kẽm|mạ\s*k[eê]m|ma\s*kẽm|ma\s+k[eê]m|\bkẽm\b|\bkem\b|zinc|galvanized'
    if re.search(_MA_KEM_PAT, sl):
        if re.search(r'\b12[.,]9\b', sl): return 'Thép mạ kẽm 12.9'
        if re.search(r'\b10[.,]9\b', sl): return 'Thép mạ kẽm 10.9'
        return 'Thép mạ kẽm 8.8'
    if re.search(r'\b12[.,]9\b',                       sl): return 'Thép 12.9'
    if re.search(r'\b10[.,]9\b',                       sl): return 'Thép 10.9'
    if re.search(r'\b8[.,]8\b',                        sl): return 'Thép 8.8'
    # Plain steel — return empty so product_matcher uses its own detection
    return ''


# ─────────────────────────────────────────────
# TEXT EXTRACTORS (Word / PDF)
# ─────────────────────────────────────────────

def _extract_text_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append('\t'.join(cells))
    return '\n'.join(parts)


def _extract_text_pdf(filepath: str) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                if tables:
                    for tbl in tables:
                        for row in tbl:
                            if row:
                                cells = [str(c or '').strip() for c in row]
                                if any(cells):
                                    parts.append('\t'.join(cells))
                else:
                    txt = page.extract_text()
                    if txt:
                        parts.append(txt)
        return '\n'.join(parts)
    except Exception:
        return ''


def _pdf_to_images(filepath: str) -> List[str]:
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(filepath, dpi=200, first_page=1, last_page=5)
        tmp = tempfile.mkdtemp()
        paths = []
        for i, img in enumerate(images):
            p = str(Path(tmp) / f'page_{i+1}.jpg')
            img.save(p, 'JPEG', quality=85)
            paths.append(p)
        return paths
    except Exception:
        return []


# ─────────────────────────────────────────────
# AI EXTRACTION PROMPT
# ─────────────────────────────────────────────

_PROMPT = """Bạn là AI phân tích tài liệu đơn hàng vật tư công nghiệp (bulong, đai ốc, đệm, vít, chốt chẻ, đinh rút...).

Hãy đọc toàn bộ nội dung và trích xuất DANH SÁCH HÀNG HÓA cần mua.

Trả về JSON array (CHỈ JSON, không có chú thích hay text thêm), mỗi phần tử:
{
  "name": "tên sản phẩm và quy cách đầy đủ, VD: Bu lông M10x30, Lục giác chìm M8x20, Đệm phẳng M12, ECU M10, Đệm vênh M10, Chốt chẻ D4x25, Đinh rút A4x10",
  "material": "chất liệu nếu có — Inox 304 / Inox 316 / Inox 201 / Thép 8.8 / Thép 10.9 / Thép 12.9 / Thép mạ kẽm 8.8 — chuỗi rỗng nếu không rõ",
  "qty": số lượng (số nguyên, mặc định 1 nếu không nêu),
  "note_in": "ghi chú nếu có, không thì chuỗi rỗng"
}

Quy tắc:
- PHẢI trích xuất TẤT CẢ các dòng hàng hóa, không bỏ sót
- Giữ nguyên ký hiệu kích thước (M8, M10x30...) trong name
- Bỏ qua: dòng tiêu đề cột, thông tin công ty/ngày tháng, dòng "Tổng"
- Nếu không tìm thấy sản phẩm nào → trả về []"""


# ─────────────────────────────────────────────
# AI EXTRACT FROM TEXT (Word / PDF)
# ─────────────────────────────────────────────

def _ai_extract_from_text(text: str) -> List[dict]:
    if not text or not text.strip():
        return []
    import anthropic
    client = anthropic.Anthropic()
    full_prompt = f"{_PROMPT}\n\nNội dung tài liệu:\n{text[:15000]}"
    try:
        resp = client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=4000,
            messages=[{'role': 'user', 'content': full_prompt}],
        )
        return _parse_ai_response(resp.content[0].text)
    except Exception as e:
        print(f'[input_parser] AI text error: {e}')
        raise RuntimeError(f'Lỗi AI khi phân tích file: {e}') from e


# ─────────────────────────────────────────────
# AI EXTRACT FROM IMAGES
# ─────────────────────────────────────────────

_EXT_MIME = {
    'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
    'png': 'image/png',  'webp': 'image/webp',
    'bmp': 'image/png',  'tiff': 'image/png', 'tif': 'image/png',
}


def _ai_extract_from_images(image_paths: List[str]) -> List[dict]:
    if not image_paths:
        return []
    import anthropic
    client = anthropic.Anthropic()

    content = []
    for img_path in image_paths[:5]:
        try:
            ext = Path(img_path).suffix.lower().lstrip('.')
            mime = _EXT_MIME.get(ext, 'image/jpeg')
            with open(img_path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode()
            content.append({
                'type': 'image',
                'source': {'type': 'base64', 'media_type': mime, 'data': b64},
            })
        except Exception:
            continue

    if not content:
        return []

    content.append({'type': 'text', 'text': _PROMPT})

    try:
        resp = client.messages.create(
            model='claude-haiku-4-5-20251001',
            max_tokens=4000,
            messages=[{'role': 'user', 'content': content}],
        )
        return _parse_ai_response(resp.content[0].text)
    except Exception as e:
        print(f'[input_parser] AI vision error: {e}')
        raise RuntimeError(f'Lỗi AI khi phân tích ảnh: {e}') from e


# ─────────────────────────────────────────────
# RESPONSE PARSER
# ─────────────────────────────────────────────

def _parse_ai_response(text: str) -> List[dict]:
    text = text.strip()
    m = re.search(r'\[.*\]', text, re.DOTALL)
    if not m:
        return []
    try:
        data = json.loads(m.group())
    except json.JSONDecodeError:
        return []

    results = []
    for item in data:
        if not isinstance(item, dict):
            continue
        name = str(item.get('name', '') or '').strip()
        if not name:
            continue
        try:
            qty = max(1, int(float(str(item.get('qty', 1) or 1))))
        except (ValueError, TypeError):
            qty = 1
        results.append({
            'name':     name,
            'material': str(item.get('material', '') or '').strip(),
            'qty':      qty,
            'note_in':  str(item.get('note_in', '') or '').strip(),
        })
    return results
